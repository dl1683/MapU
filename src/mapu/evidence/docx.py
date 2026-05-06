"""DOCX document parser using python-docx."""

from __future__ import annotations

import io

from docx import Document

from mapu.evidence.types import (
    DocumentBlob,
    ParsedDocument,
    ParsedNode,
    ParsedSpan,
)


class DocxParser:
    """Parses DOCX documents into structured nodes (paragraphs) and text spans."""

    @property
    def parser_id(self) -> str:
        return "docx_python_docx_v1"

    @property
    def supported_mime_types(self) -> frozenset[str]:
        return frozenset({
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        })

    async def parse(self, document: DocumentBlob) -> ParsedDocument:
        doc = Document(io.BytesIO(document.content))

        nodes: list[ParsedNode] = []
        spans: list[ParsedSpan] = []
        full_parts: list[str] = []
        char_offset = 0

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else "Normal"
            node_type = "heading" if style_name.startswith("Heading") else "paragraph"

            nodes.append(ParsedNode(
                node_type=node_type,
                ordinal=len(nodes),
                text=text,
                metadata={"style": style_name},
            ))

            spans.append(ParsedSpan(
                text=text,
                start_char=char_offset,
                end_char=char_offset + len(text),
                node_index=len(nodes) - 1,
            ))

            full_parts.append(text)
            char_offset += len(text) + 1

        full_text = "\n".join(full_parts)

        return ParsedDocument(
            parser_id=self.parser_id,
            nodes=tuple(nodes),
            spans=tuple(spans),
            full_text=full_text,
            metadata={
                "source_uri": document.source_uri,
                "paragraph_count": len(nodes),
            },
        )
