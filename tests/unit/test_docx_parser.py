"""Unit tests for DOCX parser."""

from __future__ import annotations

import io

import pytest
from docx import Document  # type: ignore[import-untyped]

from mapu.evidence.docx import DocxParser
from mapu.evidence.types import DocumentBlob


def _make_docx(paragraphs: list[str]) -> bytes:
    """Create a minimal DOCX with given paragraphs."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_with_heading(heading: str, body: str) -> bytes:
    """Create a DOCX with a heading and body paragraph."""
    doc = Document()
    doc.add_heading(heading, level=1)
    doc.add_paragraph(body)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxParser:
    @pytest.fixture
    def parser(self) -> DocxParser:
        return DocxParser()

    async def test_parser_id(self, parser: DocxParser) -> None:
        assert parser.parser_id == "docx_python_docx_v1"

    async def test_supported_types(self, parser: DocxParser) -> None:
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert mime in parser.supported_mime_types

    async def test_parse_simple(self, parser: DocxParser) -> None:
        docx_bytes = _make_docx(["First paragraph.", "Second paragraph."])
        blob = DocumentBlob(
            content=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_uri="test://simple.docx",
        )
        result = await parser.parse(blob)
        assert len(result.nodes) == 2
        assert result.nodes[0].text == "First paragraph."
        assert result.nodes[1].text == "Second paragraph."
        assert len(result.spans) == 2

    async def test_empty_document(self, parser: DocxParser) -> None:
        docx_bytes = _make_docx([])
        blob = DocumentBlob(
            content=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_uri="test://empty.docx",
        )
        result = await parser.parse(blob)
        assert len(result.nodes) == 0
        assert len(result.spans) == 0

    async def test_heading_detection(self, parser: DocxParser) -> None:
        docx_bytes = _make_docx_with_heading("My Title", "Body text here.")
        blob = DocumentBlob(
            content=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_uri="test://heading.docx",
        )
        result = await parser.parse(blob)
        assert any(n.node_type == "heading" for n in result.nodes)
        assert any(n.node_type == "paragraph" for n in result.nodes)

    async def test_span_offsets(self, parser: DocxParser) -> None:
        docx_bytes = _make_docx(["Hello.", "World."])
        blob = DocumentBlob(
            content=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_uri="test://offsets.docx",
        )
        result = await parser.parse(blob)
        assert result.spans[0].start_char == 0
        assert result.spans[0].end_char == 6
        assert result.spans[1].start_char == 7
        assert result.spans[1].end_char == 13

    async def test_metadata_has_paragraph_count(self, parser: DocxParser) -> None:
        docx_bytes = _make_docx(["One.", "Two.", "Three."])
        blob = DocumentBlob(
            content=docx_bytes,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            source_uri="test://meta.docx",
        )
        result = await parser.parse(blob)
        assert result.metadata["paragraph_count"] == 3
